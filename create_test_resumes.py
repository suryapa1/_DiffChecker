import os
import tempfile
import docx
from docx import Document
from docx.shared import Pt
import random

def create_test_resumes():
    """
    Create test resume files for validating the diff checker application
    """
    test_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_files")
    os.makedirs(test_dir, exist_ok=True)
    
    # Create first resume
    resume1_path = os.path.join(test_dir, "resume1.docx")
    create_resume_docx(
        resume1_path,
        name="John Smith",
        email="john.smith@example.com",
        phone="(555) 123-4567",
        summary="Experienced software developer with 5 years of experience in web development and cloud technologies.",
        skills=["Python", "JavaScript", "React", "Node.js", "AWS", "Docker", "Git", "Agile"],
        experience=[
            {
                "title": "Senior Developer",
                "company": "Tech Solutions Inc.",
                "dates": "2020 - Present",
                "description": "Lead developer for cloud-based applications. Implemented CI/CD pipelines and microservices architecture."
            },
            {
                "title": "Web Developer",
                "company": "Digital Innovations",
                "dates": "2018 - 2020",
                "description": "Developed responsive web applications using React and Node.js. Collaborated with UX designers to implement user-friendly interfaces."
            }
        ],
        education=[
            {
                "degree": "Bachelor of Science in Computer Science",
                "school": "University of Technology",
                "dates": "2014 - 2018",
                "description": "GPA: 3.8/4.0. Relevant coursework: Data Structures, Algorithms, Web Development, Database Systems."
            }
        ],
        certifications=["AWS Certified Developer", "Certified Scrum Master"]
    )
    
    # Create second resume (similar but with differences)
    resume2_path = os.path.join(test_dir, "resume2.docx")
    create_resume_docx(
        resume2_path,
        name="John Smith",
        email="johnsmith@gmail.com",  # Different email
        phone="(555) 123-4567",
        summary="Experienced software developer with 6 years of experience in web development, cloud technologies, and mobile applications.",  # Added mobile
        skills=["Python", "JavaScript", "React", "Angular", "Node.js", "AWS", "Docker", "Kubernetes", "Git", "Agile"],  # Added Angular and Kubernetes
        experience=[
            {
                "title": "Senior Developer",
                "company": "Tech Solutions Inc.",
                "dates": "2020 - Present",
                "description": "Lead developer for cloud-based applications. Implemented CI/CD pipelines and microservices architecture. Mentored junior developers and conducted code reviews."  # Added mentoring
            },
            {
                "title": "Web Developer",
                "company": "Digital Innovations",
                "dates": "2018 - 2020",
                "description": "Developed responsive web applications using React and Node.js. Collaborated with UX designers to implement user-friendly interfaces."
            },
            {
                "title": "Junior Developer",  # Added new position
                "company": "StartUp Tech",
                "dates": "2017 - 2018",
                "description": "Developed and maintained company website and internal tools."
            }
        ],
        education=[
            {
                "degree": "Bachelor of Science in Computer Science",
                "school": "University of Technology",
                "dates": "2014 - 2018",
                "description": "GPA: 3.8/4.0. Relevant coursework: Data Structures, Algorithms, Web Development, Database Systems."
            }
        ],
        certifications=["AWS Certified Developer", "Certified Scrum Master", "Google Cloud Associate Engineer"]  # Added GCP cert
    )
    
    # Create Word documents with different content
    word_doc1_path = os.path.join(test_dir, "document1.docx")
    doc1 = Document()
    doc1.add_heading('Project Proposal', 0)
    doc1.add_paragraph('This document outlines the proposal for the new software project.')
    doc1.add_heading('Objectives', level=1)
    doc1.add_paragraph('The main objectives of this project are:')
    doc1.add_paragraph('1. Develop a user-friendly interface', style='List Number')
    doc1.add_paragraph('2. Implement robust backend services', style='List Number')
    doc1.add_paragraph('3. Ensure scalability and performance', style='List Number')
    doc1.add_heading('Timeline', level=1)
    doc1.add_paragraph('The project will be completed in 3 months.')
    doc1.add_heading('Budget', level=1)
    doc1.add_paragraph('The estimated budget is $50,000.')
    doc1.save(word_doc1_path)
    
    word_doc2_path = os.path.join(test_dir, "document2.docx")
    doc2 = Document()
    doc2.add_heading('Project Proposal - Revised', 0)
    doc2.add_paragraph('This document outlines the revised proposal for the new software project.')
    doc2.add_heading('Objectives', level=1)
    doc2.add_paragraph('The main objectives of this project are:')
    doc2.add_paragraph('1. Develop a user-friendly interface with modern design', style='List Number')
    doc2.add_paragraph('2. Implement robust backend services with high availability', style='List Number')
    doc2.add_paragraph('3. Ensure scalability and performance', style='List Number')
    doc2.add_paragraph('4. Add mobile support', style='List Number')
    doc2.add_heading('Timeline', level=1)
    doc2.add_paragraph('The project will be completed in 4 months due to added requirements.')
    doc2.add_heading('Budget', level=1)
    doc2.add_paragraph('The revised estimated budget is $65,000.')
    doc2.add_heading('Team', level=1)
    doc2.add_paragraph('The project team will consist of 5 developers and 2 designers.')
    doc2.save(word_doc2_path)
    
    # Create a zip file with multiple files for folder testing
    import zipfile
    zip_path = os.path.join(test_dir, "resume_package.zip")
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        zipf.write(resume1_path, arcname="resume1.docx")
        zipf.write(word_doc1_path, arcname="document1.docx")
        
        # Add a binary file to test filtering
        bin_file_path = os.path.join(test_dir, "binary_file.bin")
        with open(bin_file_path, 'wb') as f:
            f.write(os.urandom(1024))
        zipf.write(bin_file_path, arcname="binary_file.bin")
    
    zip_path2 = os.path.join(test_dir, "resume_package2.zip")
    with zipfile.ZipFile(zip_path2, 'w') as zipf:
        zipf.write(resume2_path, arcname="resume1.docx")
        zipf.write(word_doc2_path, arcname="document1.docx")
        
        # Add a nested zip to test recursive extraction
        nested_zip_path = os.path.join(test_dir, "nested.zip")
        with zipfile.ZipFile(nested_zip_path, 'w') as nested_zipf:
            nested_zipf.write(resume1_path, arcname="original_resume.docx")
        zipf.write(nested_zip_path, arcname="nested.zip")
    
    # Clean up temporary files
    if os.path.exists(bin_file_path):
        os.remove(bin_file_path)
    if os.path.exists(nested_zip_path):
        os.remove(nested_zip_path)
    
    return test_dir

def create_resume_docx(filepath, name, email, phone, summary, skills, experience, education, certifications):
    """
    Create a Word document resume with the given information
    """
    doc = Document()
    
    # Add name as title
    title = doc.add_heading(name, 0)
    for run in title.runs:
        run.font.size = Pt(16)
    
    # Add contact info
    contact_para = doc.add_paragraph()
    contact_para.add_run(f"Email: {email} | Phone: {phone}")
    
    # Add summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    
    # Add skills
    doc.add_heading('Skills', level=1)
    skills_para = doc.add_paragraph()
    skills_para.add_run(', '.join(skills))
    
    # Add experience
    doc.add_heading('Experience', level=1)
    for job in experience:
        p = doc.add_paragraph()
        p.add_run(f"{job['title']} at {job['company']}").bold = True
        p.add_run(f" ({job['dates']})")
        doc.add_paragraph(job['description'])
    
    # Add education
    doc.add_heading('Education', level=1)
    for edu in education:
        p = doc.add_paragraph()
        p.add_run(f"{edu['degree']} - {edu['school']}").bold = True
        p.add_run(f" ({edu['dates']})")
        doc.add_paragraph(edu['description'])
    
    # Add certifications
    if certifications:
        doc.add_heading('Certifications', level=1)
        for cert in certifications:
            doc.add_paragraph(cert, style='List Bullet')
    
    doc.save(filepath)

if __name__ == "__main__":
    test_dir = create_test_resumes()
    print(f"Test files created in: {test_dir}")
