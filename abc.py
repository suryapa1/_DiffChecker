import streamlit as st
import os
import docx
import re
import base64
from io import BytesIO
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import tempfile
import shutil
import markdown
from weasyprint import HTML
import subprocess

# Set page configuration
st.set_page_config(
    page_title="Surya's Resume Generator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #2563EB;
        margin-bottom: 1rem;
    }
    .info-box {
        background-color: #EFF6FF;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #3B82F6;
        margin-bottom: 1rem;
    }
    .success-box {
        background-color: #ECFDF5;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #10B981;
        margin-bottom: 1rem;
    }
    .warning-box {
        background-color: #FFFBEB;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #F59E0B;
        margin-bottom: 1rem;
    }
    .skill-match {
        background-color: #DCFCE7;
        padding: 0.2rem 0.5rem;
        border-radius: 0.25rem;
    }
    .skill-missing {
        background-color: #FEF2F2;
        padding: 0.2rem 0.5rem;
        border-radius: 0.25rem;
    }
    .download-button {
        display: inline-block;
        padding: 0.5rem 1rem;
        background-color: #2563EB;
        color: white;
        text-decoration: none;
        border-radius: 0.25rem;
        margin-right: 0.5rem;
        margin-bottom: 0.5rem;
    }
    .download-button:hover {
        background-color: #1E40AF;
    }
    .guidance-item {
        margin-bottom: 1rem;
        padding: 0.5rem;
        border-radius: 0.25rem;
        background-color: #F3F4F6;
    }
    .guidance-title {
        font-weight: bold;
        color: #1E3A8A;
    }
    .tab-content {
        padding: 1rem;
        border: 1px solid #E5E7EB;
        border-radius: 0.5rem;
        margin-top: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state variables
if 'jd_text' not in st.session_state:
    st.session_state.jd_text = ""
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = ""
if 'generated_resume' not in st.session_state:
    st.session_state.generated_resume = ""
if 'gap_analysis' not in st.session_state:
    st.session_state.gap_analysis = ""
if 'resume_format' not in st.session_state:
    st.session_state.resume_format = "hr_focused"
if 'jd_analysis' not in st.session_state:
    st.session_state.jd_analysis = None
if 'extracted_skills' not in st.session_state:
    st.session_state.extracted_skills = []
if 'highlighted_resume' not in st.session_state:
    st.session_state.highlighted_resume = ""
if 'improvement_guidance' not in st.session_state:
    st.session_state.improvement_guidance = ""


# Load resume templates from resources
def load_resume_templates():
    templates = {}
    template_files = {
        "hr_focused": "/home/ubuntu/resume_app/resources/hr_focused_resume.md",
        "semi_technical": "/home/ubuntu/resume_app/resources/semi_technical_resume.md",
        "design_oriented": "/home/ubuntu/resume_app/resources/design_oriented_resume.md",
        "ai_architecture": "/home/ubuntu/resume_app/resources/ai_architecture_resume.md"
    }

    for format_type, file_path in template_files.items():
        try:
            with open(file_path, 'r') as f:
                templates[format_type] = f.read()
        except Exception as e:
            templates[format_type] = f"Error loading template: {str(e)}"

    return templates


# Function to extract text from DOCX
def extract_text_from_docx(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


# Function to extract text from PDF using poppler-utils
def extract_text_from_pdf(file):
    try:
        # First try using pdftotext from poppler-utils
        with tempfile.NamedTemporaryFile(delete=False, suffix='.txt') as temp_txt:
            temp_txt_path = temp_txt.name

        # Run pdftotext command
        subprocess.run(['pdftotext', file, temp_txt_path], check=True)

        # Read the extracted text
        with open(temp_txt_path, 'r') as f:
            text = f.read()

        # Clean up
        os.unlink(temp_txt_path)

        return text
    except Exception as e:
        # Fallback to PyPDF2 if poppler-utils fails
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text()
            return text
        except Exception as e2:
            return f"Error extracting PDF text: {str(e2)}"


# Function to create DOCX from text
def create_docx_from_text(text, format_type):
    doc = docx.Document()

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("SURYA NERSU")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process text by paragraphs
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            if para.startswith('## '):
                p = doc.add_paragraph()
                run = p.add_run(para[3:])
                run.bold = True
                run.font.size = Pt(14)
            elif para.startswith('### '):
                p = doc.add_paragraph()
                run = p.add_run(para[4:])
                run.bold = True
                run.font.size = Pt(12)
            elif para.startswith('- '):
                p = doc.add_paragraph(para[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(para)

    # Save to BytesIO
    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    return docx_io


# Function to create PDF from text
def create_pdf_from_text(text, format_type):
    # Convert markdown to HTML
    html_content = markdown.markdown(text)

    # Add basic styling
    styled_html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ text-align: center; font-size: 18pt; }}
            h2 {{ font-size: 16pt; margin-top: 20px; }}
            h3 {{ font-size: 14pt; }}
            p {{ font-size: 12pt; line-height: 1.5; }}
            ul {{ padding-left: 20px; }}
            li {{ margin-bottom: 5px; }}
            .center {{ text-align: center; }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    '''

    # Generate PDF to BytesIO
    pdf_io = BytesIO()
    HTML(string=styled_html).write_pdf(pdf_io)
    pdf_io.seek(0)

    return pdf_io


# Function to create download link
def get_binary_file_downloader_html(bin_file, file_label, file_extension):
    bin_file.seek(0)
    data = bin_file.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/{file_extension};base64,{b64}" download="{file_label}.{file_extension}" class="download-button">Download {file_label}</a>'
    return href


# Function to analyze JD and extract key requirements
def analyze_jd(jd_text):
    # Extract skills (simplified version)
    skills_pattern = r'(?:skills|requirements|qualifications|experience with|proficiency in|knowledge of)(?:[^\n]*)((?:[\s]*[-•][\s]*[^\n]+[\n]?)+)'
    skills_matches = re.findall(skills_pattern, jd_text.lower(), re.IGNORECASE)

    skills = []
    for match in skills_matches:
        skill_items = re.findall(r'[-•][\s]*([^\n]+)', match)
        skills.extend([item.strip() for item in skill_items if item.strip()])

    # Extract years of experience
    exp_pattern = r'(\d+\+?)\s*(?:years|yrs)(?:\s*of)?\s*(?:experience|exp)'
    exp_matches = re.findall(exp_pattern, jd_text.lower())
    experience_req = exp_matches[0] if exp_matches else "Not specified"

    # Extract education
    edu_pattern = r'(?:degree|bachelor|master|phd|bs|ms|b\.s\.|m\.s\.)(?:[^\n]*)((?:[\s]*[-•][\s]*[^\n]+[\n]?)+)'
    edu_matches = re.findall(edu_pattern, jd_text.lower(), re.IGNORECASE)

    education = []
    for match in edu_matches:
        edu_items = re.findall(r'[-•][\s]*([^\n]+)', match)
        education.extend([item.strip() for item in edu_items if item.strip()])

    # Extract key technologies
    tech_keywords = [
        "python", "java", "javascript", "c#", "c++", "ruby", "go", "rust", "php",
        "aws", "azure", "gcp", "cloud", "kubernetes", "docker", "terraform",
        "react", "angular", "vue", "node", "django", "flask", "spring", "asp.net",
        "sql", "nosql", "mongodb", "postgresql", "mysql", "oracle", "database",
        "machine learning", "ml", "ai", "artificial intelligence", "deep learning", "neural network",
        "data science", "big data", "hadoop", "spark", "kafka", "data engineering",
        "devops", "ci/cd", "jenkins", "github actions", "gitlab ci", "automation",
        "agile", "scrum", "kanban", "jira", "product management", "project management",
        "microservices", "rest api", "graphql", "serverless", "lambda", "function",
        "security", "encryption", "authentication", "authorization", "oauth",
        "mobile", "ios", "android", "react native", "flutter", "xamarin",
        "testing", "qa", "quality assurance", "unit testing", "integration testing",
        "blockchain", "ethereum", "smart contract", "web3", "cryptocurrency",
        "llm", "large language model", "gpt", "bert", "transformer", "nlp",
        "computer vision", "cv", "image processing", "object detection",
        "rag", "retrieval augmented generation", "vector database", "embedding",
        "langchain", "langgraph", "multi-agent", "agent", "orchestration",
        "fine-tuning", "qlora", "lora", "prompt engineering",
        "kubernetes", "k8s", "container", "docker", "microservice",
        "observability", "monitoring", "logging", "grafana", "prometheus",
        "mlops", "ml engineering", "model deployment", "model serving",
        "data pipeline", "etl", "elt", "data transformation", "data lake", "data warehouse",
        "vector search", "vector store", "redis", "pgvector", "pinecone", "weaviate",
        "hipaa", "gdpr", "compliance", "regulated", "pii", "phi", "redaction",
        "salesforce", "sfdc", "crm", "customer relationship management"
    ]

    technologies = []
    for keyword in tech_keywords:
        if re.search(r'\b' + re.escape(keyword) + r'\b', jd_text.lower()):
            technologies.append(keyword)

    # Extract job title
    title_patterns = [
        r'job title:?\s*([^\n]+)',
        r'position:?\s*([^\n]+)',
        r'role:?\s*([^\n]+)',
        r'^([^:]+(?:engineer|developer|architect|scientist|analyst|manager|director|lead|head|chief|vp|president|officer|specialist|consultant|advisor))[\s\n]'
    ]

    job_title = "Not specified"
    for pattern in title_patterns:
        matches = re.findall(pattern, jd_text, re.IGNORECASE | re.MULTILINE)
        if matches:
            job_title = matches[0].strip()
            break

    # Extract company name
    company_patterns = [
        r'(?:at|with|for|join)\s+([A-Z][A-Za-z0-9\s&]+)(?:[\.,]|\s+is|\s+as|\s+to)',
        r'about\s+([A-Z][A-Za-z0-9\s&]+)(?:[\.,]|\s+is|\s+as|\s+to)',
        r'([A-Z][A-Za-z0-9\s&]+)\s+is\s+(?:looking|seeking|hiring)',
    ]

    company_name = "Not specified"
    for pattern in company_patterns:
        matches = re.findall(pattern, jd_text)
        if matches:
            company_name = matches[0].strip()
            break

    return {
        "job_title": job_title,
        "company_name": company_name,
        "skills": skills,
        "technologies": technologies,
        "experience": experience_req,
        "education": education
    }


# Function to perform gap analysis
def perform_gap_analysis(jd_analysis, resume_text):
    missing_skills = []
    matching_skills = []

    # Check for skills
    for skill in jd_analysis["skills"]:
        if skill.strip() and skill.strip().lower() not in resume_text.lower():
            missing_skills.append(skill.strip())
        else:
            matching_skills.append(skill.strip())

    # Check for technologies
    missing_tech = []
    matching_tech = []
    for tech in jd_analysis["technologies"]:
        if tech.strip() and tech.strip().lower() not in resume_text.lower():
            missing_tech.append(tech.strip())
        else:
            matching_tech.append(tech.strip())

    # Calculate match percentage
    total_skills = len(jd_analysis["skills"]) + len(jd_analysis["technologies"])
    matching_count = len(matching_skills) + len(matching_tech)

    match_percentage = 0
    if total_skills > 0:
        match_percentage = round((matching_count / total_skills) * 100)

    analysis = f"""
## Gap Analysis

### Job Overview
- **Position**: {jd_analysis["job_title"]}
- **Company**: {jd_analysis["company_name"]}
- **Experience Required**: {jd_analysis["experience"]} years
- **Overall Match**: {match_percentage}% of required skills present in your resume

### Skills Match
"""

    if matching_skills:
        analysis += "The following skills from the job description are present in your resume:\n"
        for skill in matching_skills:
            analysis += f"\n- <span class='skill-match'>{skill}</span>"
    else:
        analysis += "No direct skill matches were found between the job description and your resume.\n"

    analysis += "\n\n### Skills Gap\n"

    if missing_skills:
        analysis += "The following skills mentioned in the job description may be missing or not prominently featured in your resume:\n"
        for skill in missing_skills:
            analysis += f"\n- <span class='skill-missing'>{skill}</span>"
    else:
        analysis += "No significant skill gaps identified!\n"

    analysis += "\n\n### Technology Match\n"

    if matching_tech:
        analysis += "The following technologies from the job description are present in your resume:\n"
        for tech in matching_tech:
            analysis += f"\n- <span class='skill-match'>{tech}</span>"

    analysis += "\n\n### Technology Gap\n"

    if missing_tech:
        analysis += "The following technologies mentioned in the job description may be missing or not prominently featured in your resume:\n"
        for tech in missing_tech:
            analysis += f"\n- <span class='skill-missing'>{tech}</span>"
    else:
        analysis += "No significant technology gaps identified!\n"

    return analysis, missing_skills + missing_tech, matching_skills + matching_tech


# Function to generate improvement guidance
def generate_improvement_guidance(missing_skills, matching_skills, resume_format):
    guidance = """
## Resume Improvement Guidance

### Recommended Actions
"""

    # Add specific recommendations based on missing skills
    if missing_skills:
        guidance += """
1. **Add Missing Skills**: Consider adding these missing skills to your resume if you have experience with them:
"""
        for skill in missing_skills[:5]:  # Limit to top 5 to avoid overwhelming
            guidance += f"\n- <span class='skill-missing'>{skill}</span>"

            # Add specific advice for common skills
            if "cloud" in skill.lower() or "aws" in skill.lower() or "azure" in skill.lower():
                guidance += " - Mention specific cloud services you've worked with and projects deployed"
            elif "machine learning" in skill.lower() or "ai" in skill.lower():
                guidance += """
 - Mention relevant ML projects
 - Include frameworks like TensorFlow, PyTorch, etc.
"""